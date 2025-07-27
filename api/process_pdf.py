import json
from flask import Flask, request, jsonify, send_file
from pdf2docx import Converter
from docx import Document
from .utils import allowed_file, generate_temp_path, cleanup_files
from .constants import TEMP_FOLDER
import tempfile

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB

@app.route('/process-pdf', methods=['POST'])
def process_pdf():
    # Check if file was uploaded
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': 'Only PDF files are allowed'}), 400

    # Parse replacements
    replacements = {}
    if request.form.get('replacements'):
        try:
            replacements = json.loads(request.form.get('replacements'))
        except json.JSONDecodeError:
            return jsonify({'error': 'Invalid replacements format'}), 400

    # Create temporary files
    pdf_path = generate_temp_path('.pdf')
    docx_path = generate_temp_path('.docx')
    output_pdf_path = generate_temp_path('.pdf')

    try:
        # Save uploaded file
        file.save(pdf_path)

        # Convert PDF to DOCX
        cv = Converter(pdf_path)
        cv.convert(docx_path)
        cv.close()

        # Edit DOCX if replacements provided
        if replacements:
            doc = Document(docx_path)
            
            def replace_text(runs):
                for run in runs:
                    for old_text, new_text in replacements.items():
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)

            # Replace in paragraphs
            for para in doc.paragraphs:
                replace_text(para.runs)

            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            replace_text(para.runs)

            doc.save(docx_path)

        # Convert back to PDF (using pdf2docx built-in)
        cv = Converter(pdf_path)
        cv.convert(output_pdf_path)
        cv.close()

        # Return the processed file
        return send_file(
            output_pdf_path,
            as_attachment=True,
            download_name=f"processed_{secure_filename(file.filename)}",
            mimetype='application/pdf'
        )

    except Exception as e:
        return jsonify({'error': str(e)}), 500

    finally:
        # Cleanup temporary files
        cleanup_files(pdf_path, docx_path, output_pdf_path)

# Vercel requires this for serverless functions
def handler(request):
    with app.app_context():
        if request.method == 'POST':
            return process_pdf()
        return jsonify({'error': 'Method not allowed'}), 405