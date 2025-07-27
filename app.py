import os
import uuid
import time
import json
import subprocess
import shutil
import threading
import logging
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from pdf2docx import Converter
from docx import Document
from werkzeug.utils import secure_filename

# Konfigurasi logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)

app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {'pdf'}
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# ----------- Fungsi Konversi DOCX ke PDF via LibreOffice -----------
def convert_docx_to_pdf(docx_path, output_pdf_path=None):
    if not shutil.which("soffice") and not shutil.which("libreoffice"):
        raise EnvironmentError("LibreOffice CLI tidak ditemukan. Pastikan sudah terinstall.")

    output_dir = os.path.dirname(output_pdf_path or docx_path)
    cmd = [
        "soffice",
        "--headless",
        "--convert-to", "pdf",
        "--outdir", output_dir,
        docx_path
    ]

    try:
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        generated_pdf = os.path.splitext(docx_path)[0] + ".pdf"

        if output_pdf_path and os.path.abspath(generated_pdf) != os.path.abspath(output_pdf_path):
            shutil.move(generated_pdf, output_pdf_path)

        logger.info(f"Konversi DOCX ke PDF selesai: {output_pdf_path or generated_pdf}")
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f"Gagal konversi DOCX ke PDF: {e.stderr.decode('utf-8')}")

# ----------- Fungsi Validasi File -----------
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

# ----------- Konversi PDF ke DOCX -----------
def convert_pdf_to_docx(pdf_path, docx_path):
    if not os.path.exists(pdf_path):
        raise Exception(f"File PDF tidak ditemukan: {pdf_path}")

    cv = Converter(pdf_path)
    cv.convert(docx_path)
    cv.close()

    if not os.path.exists(docx_path) or os.path.getsize(docx_path) == 0:
        raise Exception(f"File DOCX tidak valid setelah konversi: {docx_path}")

# ----------- Edit Isi DOCX -----------
def edit_docx_text(docx_path, replacements):
    if not os.path.exists(docx_path):
        raise Exception(f"File DOCX tidak ditemukan: {docx_path}")

    doc = Document(docx_path)

    def replace_in_runs(runs, replacements):
        for run in runs:
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)

    for para in doc.paragraphs:
        replace_in_runs(para.runs, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_runs(para.runs, replacements)

    doc.save(docx_path)

# ----------- Hapus File Aman -----------
def safe_remove(filepath):
    try:
        if filepath and os.path.exists(filepath):
            os.remove(filepath)
            logger.info(f"File dihapus: {filepath}")
    except Exception as e:
        logger.warning(f"Gagal hapus file {filepath}: {e}")

# ----------- Route Utama -----------
@app.route('/process-pdf', methods=['POST'])
def process_pdf():
    if 'file' not in request.files:
        return jsonify({'error': 'File tidak ditemukan'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Nama file kosong'}), 400

    if not allowed_file(file.filename):
        return jsonify({'error': 'Hanya file PDF yang diperbolehkan'}), 400

    pdf_input = None
    docx_temp = None
    pdf_output = None

    try:
        replacements = {}
        if request.form.get('replacements'):
            try:
                replacements = json.loads(request.form.get('replacements'))
            except json.JSONDecodeError:
                return jsonify({'error': 'Format JSON tidak valid'}), 400

        file_id = str(uuid.uuid4())
        original_filename = secure_filename(file.filename)
        pdf_input = os.path.join(app.config['UPLOAD_FOLDER'], f"{file_id}_{original_filename}")
        docx_temp = os.path.join(app.config['UPLOAD_FOLDER'], f"{file_id}_temp.docx")
        pdf_output = os.path.join(app.config['UPLOAD_FOLDER'], f"{file_id}_modified.pdf")

        file.save(pdf_input)
        logger.info(f"File disimpan: {pdf_input}")

        convert_pdf_to_docx(pdf_input, docx_temp)
        edit_docx_text(docx_temp, replacements)
        convert_docx_to_pdf(docx_temp, pdf_output)

        return send_file(
            pdf_output,
            as_attachment=True,
            download_name=f"modified_{original_filename}",
            mimetype='application/pdf'
        )

    except Exception as e:
        logger.error(f"Proses gagal: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500

    finally:
        def cleanup():
            time.sleep(2)
            safe_remove(pdf_input)
            safe_remove(docx_temp)
            safe_remove(pdf_output)
        threading.Thread(target=cleanup).start()

# ----------- Jalankan Server -----------
if __name__ == '__main__':
    app.run(debug=True, threaded=True)
